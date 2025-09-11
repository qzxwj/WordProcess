from docx import Document
import re
from pathlib import Path

# 匹配全/半角括号中的单个大写字母（A-D）
PAT = re.compile(r'[（(][A-D]+[)）]')

def clean_paragraph_runs(par):
    full_text = ''.join(run.text for run in par.runs)
    # 如果段落中没有类似（B），直接跳过
    if not PAT.search(full_text):
        return

    # 合并所有 run 到一个，保留第一个 run 的样式
    new_text = PAT.sub('', full_text)
    first_run = par.runs[0]
    first_run.text = new_text

    # 删除其他 run
    for run in par.runs[1:]:
        run.clear()  # 清除内容，保留样式结构（不建议直接删除 Element，可能影响格式）

def clean_table(table):
    for row in table.rows:
        for cell in row.cells:
            for par in cell.paragraphs:
                clean_paragraph_runs(par)

def clean_docx(in_path, out_path):
    doc = Document(in_path)
    for p in doc.paragraphs:
        clean_paragraph_runs(p)
    for tbl in doc.tables:
        clean_table(tbl)
    doc.save(out_path)

if __name__ == "__main__":
    in_file = r"D:\我的\大三\毛概\李晨洋\毛概学习通题库.docx"
    out_file = Path(in_file).with_name("毛概学习通题库_去括号答案.docx")
    clean_docx(in_file, out_file)
    print("Done ->", out_file)
