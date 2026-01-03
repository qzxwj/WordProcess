from docx import Document

# ====== 参数区 ======
input_docx = r"D:\我的\大三\大三上\毛概\徐炜杰\毛概学习通396题（无答案）.docx"      # 原题库文件
output_docx = r"D:\我的\大三\大三上\毛概\徐炜杰\output.docx"    # 去答案后的文件

# 需要删除的关键词（出现任意一个就删该段）
ANSWER_KEYWORDS = [
    "正确答案",
    "我的答案",
    "得分"
]

# ====== 处理逻辑 ======
def remove_answer_paragraphs(input_path, output_path):
    doc = Document(input_path)
    new_doc = Document()

    for para in doc.paragraphs:
        text = para.text.strip()

        # 判断是否是答案相关段落
        if any(keyword in text for keyword in ANSWER_KEYWORDS):
            continue  # 跳过（不写入新文档）

        # 否则原样写入
        new_para = new_doc.add_paragraph()
        new_para.paragraph_format.left_indent = para.paragraph_format.left_indent
        new_para.paragraph_format.first_line_indent = para.paragraph_format.first_line_indent

        for run in para.runs:
            new_run = new_para.add_run(run.text)
            new_run.bold = run.bold
            new_run.italic = run.italic
            new_run.underline = run.underline
            new_run.font.size = run.font.size
            new_run.font.name = run.font.name

    new_doc.save(output_path)
    print(f"处理完成，结果已保存为：{output_path}")


if __name__ == "__main__":
    remove_answer_paragraphs(input_docx, output_docx)
