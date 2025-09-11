from docx import Document
import re
from pathlib import Path

# 1. 匹配括号内的选项字母（支持多个字母）
PAT_BRACKET = re.compile(r'[（(][A-D]+[)）]')

# 2. 匹配整行“正确答案： A :xxx ;”格式
PAT_ANSWER_LINE = re.compile(r'^答案：\s*[A-D]\s*$', re.UNICODE)

def clean_paragraph_runs(par):
    full_text = ''.join(run.text for run in par.runs)

    # 删除答案行：整段直接清空
    if PAT_ANSWER_LINE.match(full_text.strip()):
        par.clear()  # 整段直接清除（包括 run）
        return

    # 删除括号内选项
    if PAT_BRACKET.search(full_text):
        new_text = PAT_BRACKET.sub('', full_text)
        first_run = par.runs[0]
        first_run.text = new_text
        for run in par.runs[1:]:
            run.clear()

def clean_table(table):
    for row in table.rows:
        for cell in row.cells:
            for par in cell.paragraphs:
                clean_paragraph_runs(par)

def clean_docx(in_path, out_path):
    doc = Document(in_path)
    for p in doc.paragraphs:
        clean_paragraph_runs(p)
    for tbl in doc.tables:
        clean_table(tbl)
    doc.save(out_path)

if __name__ == "__main__":
    in_file = r"D:\我的\大三\毛概\李宇航\刷题题库.docx"
    out_file = Path(in_file).with_name("刷题题库（无答案）.docx")
    clean_docx(in_file, out_file)
    print("✅ 清理完成，输出文件：", out_file)
